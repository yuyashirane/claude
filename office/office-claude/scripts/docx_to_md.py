"""設計書 docx → Markdown 変換スクリプト

用途
----
office-claude プロジェクトの全体設計書 docx を Markdown に変換する。
段落（Heading 1〜6 のスタイル含む）と表を、本文の出現順を保って GFM Markdown 化する。
将来 docx 改訂版が来たときに再変換できるよう保持している。

使い方
------
    python office/office-claude/scripts/docx_to_md.py

末尾の SRC / DST の Path を改訂版に合わせて差し替えてから実行する。
処理後に標準出力へ統計（段落数・見出し数・表数・画像参照数）を出す。

制約・既知の挙動
----------------
- 書式情報は保持しない（太字・斜体・色・フォント・取り消し線など）
- 箇条書き(bullet/numbered list)は通常段落として出力される
  （Markdown の `-` / `1.` 記法には変換しない）
- セル内改行は `<br>`、`|` は `\\|` でエスケープ
- 画像は [画像: <filename>] プレースホルダに置換し、末尾レポートに一覧出力
- フィールドコード等 `paragraph.text` で取得できない要素は欠落する場合あり
- 連続する空段落は1行に圧縮する

作成経緯
--------
2026-05-05、`◆office-claude-design-v2.2.docx` を Markdown に変換するために作成。
原本docxには §9.2 の H3 見出し 4箇所に文字欠落があったため、変換後の md 側で
手修正している（補修履歴は docs/design/README.md 参照）。
"""
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

SRC = Path(r"C:\Users\yuya_\claude\office\office-claude\docs\design\raw\◆office-claude-design-v2.2.docx")
DST = Path(r"C:\Users\yuya_\claude\office\office-claude\docs\design\office-claude-design-v2_2.md")


def heading_level(style_name: str):
    if not style_name:
        return None
    m = re.match(r"Heading\s*(\d+)", style_name)
    if m:
        return int(m.group(1))
    m = re.match(r"見出し\s*(\d+)", style_name)
    if m:
        return int(m.group(1))
    if style_name.lower() == "title":
        return 1
    return None


def paragraph_to_md(p, image_log):
    text_parts = []
    for run in p.runs:
        # 画像チェック
        for drawing in run._element.findall('.//' + qn('w:drawing')):
            blips = drawing.findall('.//' + qn('a:blip'))
            for blip in blips:
                rid = blip.get(qn('r:embed'))
                if rid and rid in p.part.related_parts:
                    img_part = p.part.related_parts[rid]
                    fname = Path(img_part.partname).name
                    image_log.append(fname)
                    text_parts.append(f"[画像: {fname}]")
        text_parts.append(run.text or "")

    text = "".join(text_parts)
    lvl = heading_level(p.style.name if p.style else None)
    if lvl:
        text = text.strip()
        if not text:
            return ""
        lvl = max(1, min(lvl, 6))
        return f"{'#' * lvl} {text}"
    return text


def cell_to_md(cell, image_log):
    lines = []
    for p in cell.paragraphs:
        line = paragraph_to_md(p, image_log)
        # セル内の見出し記号は通常テキストにフォールバック
        line = re.sub(r"^#+\s*", "", line)
        if line.strip():
            lines.append(line.strip())
    txt = "<br>".join(lines)
    return txt.replace("|", "\\|")


def table_to_md(table, image_log):
    if not table.rows:
        return ""
    rows_data = []
    for row in table.rows:
        cells = [cell_to_md(c, image_log) for c in row.cells]
        rows_data.append(cells)

    width = max(len(r) for r in rows_data)
    rows_data = [r + [""] * (width - len(r)) for r in rows_data]

    out = []
    header = rows_data[0]
    out.append("| " + " | ".join(header) + " |")
    out.append("| " + " | ".join(["---"] * width) + " |")
    for r in rows_data[1:]:
        out.append("| " + " | ".join(r) + " |")
    return "\n".join(out)


def iter_block_items(parent):
    """本文中の paragraph と table を出現順に yield"""
    body = parent.element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag == qn('w:p'):
            yield ('p', child)
        elif tag == qn('w:tbl'):
            yield ('tbl', child)


def main():
    doc = Document(str(SRC))

    # 高速ルックアップ用に paragraph / table を辞書化
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    image_log = []
    md_parts = []
    para_count = 0
    heading_count = 0
    table_count = 0
    blank_paras = 0

    for kind, el in iter_block_items(doc):
        if kind == 'p':
            p = para_map.get(el)
            if p is None:
                continue
            para_count += 1
            line = paragraph_to_md(p, image_log)
            if line.startswith("#"):
                heading_count += 1
            if not line.strip():
                blank_paras += 1
                md_parts.append("")
            else:
                md_parts.append(line)
        elif kind == 'tbl':
            t = table_map.get(el)
            if t is None:
                continue
            table_count += 1
            md_parts.append("")
            md_parts.append(table_to_md(t, image_log))
            md_parts.append("")

    # 連続する空行をまとめる
    out_lines = []
    prev_blank = False
    for line in md_parts:
        if not line.strip():
            if not prev_blank:
                out_lines.append("")
            prev_blank = True
        else:
            out_lines.append(line)
            prev_blank = False

    DST.write_text("\n".join(out_lines).rstrip() + "\n", encoding="utf-8")

    print(f"OUT: {DST}")
    print(f"paragraphs_total: {para_count}")
    print(f"  headings: {heading_count}")
    print(f"  blank_paras: {blank_paras}")
    print(f"tables: {table_count}")
    print(f"images_referenced: {len(image_log)}")
    if image_log:
        print("image_files:")
        for f in sorted(set(image_log)):
            print(f"  - {f} (refs={image_log.count(f)})")


if __name__ == "__main__":
    main()
